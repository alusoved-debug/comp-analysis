"""
Reliability Comparison Word Report Generator  (English)
Produces a formal engineering report containing:
  - MIL-STD-217F Notice 2 analysis table
  - Manufacturer FIT-database analysis table
  - Side-by-side comparison table
  - Pareto chart of top failure contributors
  - Alternative component recommendations (top-5 by λ_p_total)
  - Summary & Conclusions
"""

import io as _io
import math
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text  import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns    import qn
from docx.oxml       import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ═══════════════════════════════════════════════════════════════════════════
#  Alternative-component database (keyed by mfr_category)
# ═══════════════════════════════════════════════════════════════════════════

ALTERNATIVES_DB: dict[str, dict] = {
    'IC_FPGA': {
        'reason': 'High gate count, large die, complex routing — dominant FIT contributor',
        'alternatives': [
            {
                'part':    'Lattice ECP5 / MachXO3L',
                'fit_25c': 500,
                'url':     'https://www.latticesemi.com/en/Support/QualityAndReliability',
                'notes':   'Flash-based config, lower gate count variant, proven industrial record',
            },
            {
                'part':    'Microchip PolarFire (MPF100T–MPF300T)',
                'fit_25c': 420,
                'url':     'https://www.microchip.com/reliabilityreport/#/',
                'notes':   'Flash FPGA, non-volatile config, low power, published qualification data',
            },
        ],
    },
    'IC_PROCESSOR': {
        'reason': 'Large SoC with GPU / ISP / PCIe subsystems — very high internal gate count',
        'alternatives': [
            {
                'part':    'NXP i.MX 8M Plus – Automotive MPU',
                'fit_25c': 260,
                'url':     'https://www.nxp.com/company/about-nxp/quality/reliability-data:RELIABILITY-DATA',
                'notes':   'ISO 26262, AEC-Q100 Grade 2, published FIT data, 16 nm FinFET',
            },
            {
                'part':    'STM32MP157 / STM32MP135 – MPU',
                'fit_25c': 185,
                'url':     'https://www.st.com/content/st_com/en/support/quality-reliability.html',
                'notes':   'Cortex-A7 + Cortex-M4, lower gate count, published qualification report',
            },
        ],
    },
    'IC_DCDC_MODULE': {
        'reason': 'µModule integrates inductor, FETs and controller in one LGA — high internal count',
        'alternatives': [
            {
                'part':    'TI TPS543x / TPS62xxx – Discrete DC-DC',
                'fit_25c': 52,
                'url':     'https://www.ti.com/quality/docs/estimator.tsp',
                'notes':   'Separate controller + external FETs + inductor; much lower FIT per function',
            },
            {
                'part':    'Renesas ISL8xx / RAA2xxxxx – Power Module',
                'fit_25c': 165,
                'url':     'https://www.renesas.com/en/support/document-search',
                'notes':   'Alternative µModule family with competitive qualification data',
            },
        ],
    },
    'IC_DESERIALIZER': {
        'reason': 'High-speed GMSL2 SerDes with many analog/digital blocks',
        'alternatives': [
            {
                'part':    'TI DS90UB954-Q1 – FPD-Link III Deserializer',
                'fit_25c': 115,
                'url':     'https://www.ti.com/quality/docs/estimator.tsp',
                'notes':   'AEC-Q100 qualified, established automotive deployment, lower complexity than GMSL2',
            },
            {
                'part':    'Renesas R-Car Companion – MIPI Bridge',
                'fit_25c': 105,
                'url':     'https://www.renesas.com/en/support/document-search',
                'notes':   'Purpose-built for camera interface; reduced feature set = lower FIT',
            },
        ],
    },
    'IC_DIGITAL_COMPLEX': {
        'reason': 'High gate-count Ethernet PHY / switch with multiple MAC-PHY blocks',
        'alternatives': [
            {
                'part':    'Microchip KSZ9031 / KSZ8081 – Ethernet PHY',
                'fit_25c': 78,
                'url':     'https://www.microchip.com/reliabilityreport/#/',
                'notes':   'Published reliability data, automotive-grade options, mature 28 nm process',
            },
            {
                'part':    'Nexperia SJA1105 – Automotive Ethernet Switch',
                'fit_25c': 85,
                'url':     'https://www.nexperia.com/quality/reliability/',
                'notes':   'TSN-capable, AEC-Q100, lower complexity with published automotive qualification',
            },
        ],
    },
    'IC_LINEAR_REG': {
        'reason': 'Complex LDO architecture with multiple protection and noise-filter circuits',
        'alternatives': [
            {
                'part':    'TI TPS7A47 / TPS7A74 – Ultra-low-noise LDO',
                'fit_25c': 32,
                'url':     'https://www.ti.com/quality/docs/estimator.tsp',
                'notes':   'Simpler pass element topology, very low noise, widely qualified',
            },
            {
                'part':    'Microchip MCP1700 / TC1014 – Low-Iq LDO',
                'fit_25c': 26,
                'url':     'https://www.microchip.com/reliabilityreport/#/',
                'notes':   'Ultra-low Iq, robust to output-capacitor ESR, proven reliability history',
            },
        ],
    },
    'IC_SUPERVISOR': {
        'reason': 'Multi-channel sequencer has many internal comparators, timers and MUXes',
        'alternatives': [
            {
                'part':    'TI TPS3xxx / Maxim MAX6xxx – Simple Supervisor',
                'fit_25c': 20,
                'url':     'https://www.ti.com/quality/docs/estimator.tsp',
                'notes':   'Single-rail supervisor, minimal gate count, many options with published MTBF',
            },
        ],
    },
    'IC_DRAM_LPDDR4': {
        'reason': 'Large DRAM die with complex refresh, ZQ calibration and power-management logic',
        'alternatives': [
            {
                'part':    'Micron LPDDR4X – Optimised process node',
                'fit_25c': 65,
                'url':     'https://reliabilityanalyticstoolkit.appspot.com',
                'notes':   'Updated 10 nm-class process, improved FIT vs. older LPDDR4 nodes',
            },
        ],
    },
    'IC_FLASH_NOR': {
        'reason': 'Floating-gate cells degrade over erase/program cycles (wear-out contribution)',
        'alternatives': [
            {
                'part':    'ISSI IS25LPxxx / IS25WPxxx – Enhanced NOR Flash',
                'fit_25c': 45,
                'url':     'https://www.issi.com/US/quality-reliability.shtml',
                'notes':   '100 k erase cycles, 20-yr data retention at 85 °C, lower FIT than older designs',
            },
            {
                'part':    'Winbond W25Qxxx – NOR Flash',
                'fit_25c': 50,
                'url':     'https://reliabilityanalyticstoolkit.appspot.com',
                'notes':   'Wide operating voltage, AEC-Q100 automotive grade available',
            },
        ],
    },
    'IC_FRAM': {
        'reason': 'Thin-film piezo layer has finite endurance under repeated write stress',
        'alternatives': [
            {
                'part':    'Cypress / Infineon FM25Vxx – FRAM (newer process)',
                'fit_25c': 28,
                'url':     'https://reliabilityanalyticstoolkit.appspot.com',
                'notes':   'Updated 40 nm ferroelectric process; 10¹⁴ endurance cycles; improved retention',
            },
        ],
    },
    'TRANSFORMER': {
        'reason': 'Wound core susceptible to humidity ingress and vibration-induced fatigue',
        'alternatives': [
            {
                'part':    'Würth Elektronik WE-LAN 749010011A – LAN Magnetics',
                'fit_25c': 32,
                'url':     'https://www.we-online.com/en/components/products/quality',
                'notes':   'SMT encapsulated winding, sealed against humidity, better vibration resistance',
            },
        ],
    },
    'CONNECTOR': {
        'reason': 'Mechanical wear and contact oxidation / fretting corrosion over mating cycles',
        'alternatives': [
            {
                'part':    'Samtec SEARAY / Harwin Gecko – High-reliability BTB',
                'fit_25c': 16,
                'url':     'https://reliabilityanalyticstoolkit.appspot.com',
                'notes':   'Gold contacts (min 0.38 µm Au), 500+ cycle rated, better fretting resistance',
            },
        ],
    },
    'CAPACITOR_ELECTROLYTIC': {
        'reason': 'Liquid electrolyte evaporation limits lifetime, especially at elevated temperature',
        'alternatives': [
            {
                'part':    'Panasonic OS-CON / KEMET A700 – Polymer Electrolytic',
                'fit_25c': 18,
                'url':     'https://reliabilityanalyticstoolkit.appspot.com',
                'notes':   'Solid conductive polymer, no electrolyte evaporation, 3× lifetime at 85 °C',
            },
        ],
    },
    'RESISTOR_WIREWOUND': {
        'reason': 'Wound element susceptible to vibration-induced open circuits',
        'alternatives': [
            {
                'part':    'Vishay WSL / Dale CSR – Thick-film chip resistor',
                'fit_25c': 8,
                'url':     'https://www.vishay.com/search/?searchChoice=part&query=',
                'notes':   'No wound element, all-SMD, lower vibration sensitivity, same power rating available',
            },
        ],
    },
    'TRANSISTOR_MOSFET': {
        'reason': 'Gate-oxide and hot-carrier degradation under continuous switching stress',
        'alternatives': [
            {
                'part':    'Infineon OptiMOS 5 / Vishay SiSxx – Next-gen Power MOSFET',
                'fit_25c': 14,
                'url':     'https://www.vishay.com/search/?searchChoice=part&query=',
                'notes':   'Latest generation; lower Rds(on)×Qg figure; lower junction temperature at same I_D',
            },
        ],
    },
    'IC_DCDC_CONTROLLER': {
        'reason': 'Complex PWM controller with many analog blocks',
        'alternatives': [
            {
                'part':    'TI LM43xxx / TPS54xxx – Simple Synchronous Buck',
                'fit_25c': 50,
                'url':     'https://www.ti.com/quality/docs/estimator.tsp',
                'notes':   'Fewer protection features, simpler layout, lower FIT for non-critical rails',
            },
        ],
    },
}


# ═══════════════════════════════════════════════════════════════════════════
#  Low-level XML / style helpers  (LTR, English)
# ═══════════════════════════════════════════════════════════════════════════

def _make_elem(tag, attribs=None):
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(k, v)
    return el


def _add_para(parent, text, bold=False, size=11, color=None,
              font='Arial', italic=False, align='left'):
    para = parent.add_paragraph()
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def _heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    para  = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name      = 'Arial'
        run.font.size      = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)


def _set_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _cell(cell, text, bold=False, size=9, align='center',
          color_rgb=None, italic=False):
    para = cell.paragraphs[0]
    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                      if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    run.font.name  = 'Arial'
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def _cell_link(cell, text, url, size=8):
    """Table cell containing a clickable hyperlink (blue, underlined)."""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not url:
        run = para.add_run(text or '—')
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        return
    try:
        r_id = cell.part.relate_to(url, RT.HYPERLINK, is_external=True)
    except Exception:
        run = para.add_run(text or url)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        return
    hl  = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    wr  = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, attrib in [
        ('w:rFonts', {qn('w:ascii'): 'Arial', qn('w:hAnsi'): 'Arial'}),
        ('w:sz',     {qn('w:val'): str(size * 2)}),
        ('w:color',  {qn('w:val'): '0563C1'}),
        ('w:u',      {qn('w:val'): 'single'}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrib.items():
            el.set(k, v)
        rPr.append(el)
    wr.append(rPr)
    t = OxmlElement('w:t')
    t.text = text or url
    wr.append(t)
    hl.append(wr)
    para._p.append(hl)


def _para_hyperlink(para, text, url, size=10):
    """Append a hyperlink run to an existing paragraph."""
    if not url:
        r = para.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        return
    try:
        r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    except Exception:
        r = para.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        return
    hl  = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    wr  = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, attrib in [
        ('w:rFonts', {qn('w:ascii'): 'Arial', qn('w:hAnsi'): 'Arial'}),
        ('w:sz',     {qn('w:val'): str(size * 2)}),
        ('w:color',  {qn('w:val'): '0563C1'}),
        ('w:u',      {qn('w:val'): 'single'}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrib.items():
            el.set(k, v)
        rPr.append(el)
    wr.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    wr.append(t)
    hl.append(wr)
    para._p.append(hl)


def _fmt(val, precision=5):
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return '—'
    if isinstance(val, float):
        return f'{val:.{precision}f}'
    return str(val)


def _compute_summary(results, key_total):
    total  = sum(r.get(key_total, 0) for r in results)
    mtbf_h = (1e6 / total) if total > 0 else 999_999_999
    return {
        'total_lambda': total,
        'fit_rate':     total * 1000,
        'mtbf_hours':   mtbf_h,
        'mtbf_years':   min(mtbf_h / 8760, 999_999),
    }


# ═══════════════════════════════════════════════════════════════════════════
#  Pareto chart
# ═══════════════════════════════════════════════════════════════════════════

def _generate_pareto_chart(results):
    """Return PNG bytes of a Pareto chart, or None if matplotlib unavailable."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    sorted_r = sorted(results,
                      key=lambda x: x.get('mfr_lambda_total', 0),
                      reverse=True)
    top_n    = min(20, len(sorted_r))
    sorted_r = sorted_r[:top_n]

    labels = [
        f"{r['ref_des']}\n{r.get('mfr_category', '?')[:12]}"
        for r in sorted_r
    ]
    values = [r.get('mfr_lambda_total', 0) for r in sorted_r]
    total  = sum(values)
    if total == 0:
        return None

    cumulative = [sum(values[:i + 1]) / total * 100 for i in range(top_n)]

    fig, ax1 = plt.subplots(figsize=(14, 6))
    colors   = ['#C0392B'] * min(5, top_n) + ['#1F3A6E'] * max(0, top_n - 5)
    ax1.bar(range(top_n), values, color=colors, alpha=0.85,
            edgecolor='white', linewidth=0.5)
    ax1.set_xlabel('Component  (sorted by total failure-rate contribution)', fontsize=10)
    ax1.set_ylabel('λ_p Total  (failures / 10⁶ hr)', color='#1F3A6E', fontsize=10)
    ax1.tick_params(axis='y', labelcolor='#1F3A6E')
    ax1.set_xticks(range(top_n))
    ax1.set_xticklabels(labels, rotation=45, ha='right', fontsize=7)

    ax2 = ax1.twinx()
    ax2.plot(range(top_n), cumulative, 'o-', color='#E67E22',
             linewidth=2, markersize=5, zorder=5)
    ax2.set_ylabel('Cumulative Contribution (%)', color='#E67E22', fontsize=10)
    ax2.tick_params(axis='y', labelcolor='#E67E22')
    ax2.set_ylim(0, 110)
    ax2.axhline(y=80, color='#E67E22', linestyle='--', alpha=0.45, linewidth=1)
    ax2.text(top_n * 0.75, 82, '80 % line', color='#E67E22', fontsize=8)

    plt.title(
        'Pareto Chart – Component Failure Rate Contribution  (Manufacturer FIT Data)',
        fontsize=11, fontweight='bold', pad=12
    )
    plt.tight_layout()

    buf = _io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════
#  Alternatives analysis
# ═══════════════════════════════════════════════════════════════════════════

def _get_top5_alternatives(results):
    """
    Identify the 5 highest-λ components (manufacturer data) and look up
    recommended alternatives from ALTERNATIVES_DB.
    Returns a dict with per-component items plus aggregated MTBF projections.
    """
    ranked = sorted(results,
                    key=lambda r: r.get('mfr_lambda_total', 0),
                    reverse=True)[:5]

    current_total    = sum(r.get('mfr_lambda_total', 0) for r in results)
    cumulative_saved = 0.0
    items            = []

    for r in ranked:
        cat    = r.get('mfr_category', '')
        qty    = r.get('quantity', 1)
        cur_lt = r.get('mfr_lambda_total', 0.0)
        entry  = ALTERNATIVES_DB.get(cat, {})
        alts   = entry.get('alternatives', [])

        if alts:
            best   = min(alts, key=lambda a: a['fit_25c'])
            new_lp = best['fit_25c'] / 1000.0
            new_lt = new_lp * qty
            saved  = max(0.0, cur_lt - new_lt)
        else:
            best   = None
            new_lt = cur_lt
            saved  = 0.0

        cumulative_saved += saved
        items.append({
            'component':    r,
            'reason':       entry.get('reason', 'High complexity / elevated FIT rate'),
            'alternatives': alts,
            'best_alt':     best,
            'cur_lt':       cur_lt,
            'new_lt':       new_lt,
            'saved':        saved,
        })

    new_total  = max(0.0, current_total - cumulative_saved)
    cur_mtbf_h = (1e6 / current_total) if current_total > 0 else float('inf')
    new_mtbf_h = (1e6 / new_total)     if new_total     > 0 else float('inf')

    return {
        'items':              items,
        'current_total':      current_total,
        'new_total':          new_total,
        'current_mtbf_h':     cur_mtbf_h,
        'new_mtbf_h':         new_mtbf_h,
        'current_mtbf_y':     cur_mtbf_h / 8760,
        'new_mtbf_y':         min(new_mtbf_h / 8760, 999_999),
        'improvement_factor': (new_mtbf_h / cur_mtbf_h) if cur_mtbf_h > 0 else 1.0,
    }


# ═══════════════════════════════════════════════════════════════════════════
#  Table builders
# ═══════════════════════════════════════════════════════════════════════════

def _make_table(doc, headers, widths_cm):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(widths_cm):
        for c in tbl.columns[i].cells:
            c.width = Cm(w)
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_shading(hdr_cells[i], '1F3A6E')
        _cell(hdr_cells[i], h, bold=True, size=9, color_rgb=(255, 255, 255))
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return tbl


def _build_mil217_table(doc, results):
    headers = [
        'Description / Part#', 'Manufacturer', 'Type', 'Qty',
        'π_T', 'π_Q', 'π_E', 'Model',
        'λ_p Unit\n(fail/10⁶hr)',
        'λ_p Total\n(fail/10⁶hr)',
    ]
    widths = [4.5, 2.8, 1.8, 0.9, 1.2, 1.2, 1.2, 2.8, 2.5, 2.5]
    tbl    = _make_table(doc, headers, widths)

    total = 0.0
    for idx, r in enumerate(results):
        d    = r.get('mil217_details', {})
        qty  = r['quantity']
        lp   = r.get('mil217_lambda_p',    0.0)
        lt   = r.get('mil217_lambda_total', 0.0)
        total += lt
        bg    = 'EEF2F7' if idx % 2 == 0 else 'FFFFFF'
        vals  = [
            f"{r['description'][:35]}\n{r['part_number'][:25]}",
            r['manufacturer'][:25],
            r['comp_type'],
            str(qty),
            _fmt(d.get('piT'), 3),
            _fmt(d.get('piQ'), 2),
            _fmt(d.get('piE'), 2),
            d.get('model', '')[:25].replace('MIL-217F ', ''),
            _fmt(lp, 5),
            _fmt(lt, 5),
        ]
        row = tbl.add_row()
        for i, v in enumerate(vals):
            _set_shading(row.cells[i], bg)
            _cell(row.cells[i], v, size=8)

    tr = tbl.add_row()
    for i in range(10):
        _set_shading(tr.cells[i], 'D4E1F7')
    _cell(tr.cells[0], 'System Total', bold=True, size=9, align='left')
    _cell(tr.cells[9], _fmt(total, 5), bold=True, size=9)


def _build_mfr_table(doc, results):
    headers = [
        'Description / Part#', 'Manufacturer', 'Category',
        'Qty', 'FIT@25°C', 'FIT@T', 'Eₐ [eV]',
        'Data Source (link)',
        'λ_p Unit\n(fail/10⁶hr)',
        'λ_p Total\n(fail/10⁶hr)',
    ]
    widths = [4.5, 2.5, 2.2, 0.9, 1.3, 1.3, 1.0, 3.5, 2.3, 2.3]
    tbl    = _make_table(doc, headers, widths)

    total = 0.0
    for idx, r in enumerate(results):
        qty      = r['quantity']
        lp       = r.get('mfr_lambda_p',    0.0)
        lt       = r.get('mfr_lambda_total', 0.0)
        total   += lt
        src_text = r.get('mfr_source', '')[:40]
        src_url  = r.get('mfr_url',    '')
        bg       = 'EEF2F7' if idx % 2 == 0 else 'FFFFFF'

        vals = [
            f"{r['description'][:35]}\n{r['part_number'][:25]}",
            r['manufacturer'][:25],
            r.get('mfr_category', '')[:20],
            str(qty),
            _fmt(r.get('mfr_fit_25c', 0), 1),
            _fmt(r.get('mfr_fit_at_t', 0), 1),
            '—',        # Ea shown in methodology section
            None,       # rendered as hyperlink below
            _fmt(lp, 5),
            _fmt(lt, 5),
        ]
        row = tbl.add_row()
        for i, v in enumerate(vals):
            _set_shading(row.cells[i], bg)
            if i == 7:
                _cell_link(row.cells[i], src_text, src_url, size=8)
            else:
                _cell(row.cells[i], v, size=8)

    tr = tbl.add_row()
    for i in range(10):
        _set_shading(tr.cells[i], 'D4E1F7')
    _cell(tr.cells[0], 'System Total', bold=True, size=9, align='left')
    _cell(tr.cells[9], _fmt(total, 5), bold=True, size=9)


def _build_comparison_table(doc, results):
    headers = [
        'Description / Part#', 'Type', 'Qty',
        'λ_p MIL-217F\n(unit)',
        'λ_p Mfr FIT\n(unit)',
        'Ratio\nMIL/Mfr',
        'λ_total MIL-217F',
        'λ_total Mfr FIT',
        'Diff %',
    ]
    widths  = [4.8, 1.8, 0.9, 2.3, 2.3, 1.4, 2.3, 2.3, 1.5]
    tbl     = _make_table(doc, headers, widths)

    total_mil = 0.0
    total_mfr = 0.0

    for idx, r in enumerate(results):
        qty   = r['quantity']
        mil_p = r.get('mil217_lambda_p',    0.0)
        mfr_p = r.get('mfr_lambda_p',       0.0)
        mil_t = r.get('mil217_lambda_total', 0.0)
        mfr_t = r.get('mfr_lambda_total',   0.0)
        total_mil += mil_t
        total_mfr += mfr_t

        ratio = r.get('ratio_mil_to_mfr')
        pct   = ((mil_t - mfr_t) / mfr_t * 100) if mfr_t > 0 else None

        ratio_str = f'{ratio:.2f}×' if ratio is not None else '—'
        pct_str   = f'{pct:+.1f}%'  if pct   is not None else '—'

        if ratio is not None and ratio > 5:
            bg = 'FFE0B2'
        elif ratio is not None and ratio > 2:
            bg = 'FFFDE7'
        elif idx % 2 == 0:
            bg = 'EEF2F7'
        else:
            bg = 'FFFFFF'

        vals = [
            f"{r['description'][:35]}\n{r['part_number'][:25]}",
            r['comp_type'],
            str(qty),
            _fmt(mil_p, 5),
            _fmt(mfr_p, 5),
            ratio_str,
            _fmt(mil_t, 5),
            _fmt(mfr_t, 5),
            pct_str,
        ]
        row = tbl.add_row()
        for i, v in enumerate(vals):
            _set_shading(row.cells[i], bg)
            _cell(row.cells[i], v, size=8)

    pct_total = ((total_mil - total_mfr) / total_mfr * 100) if total_mfr > 0 else None
    ratio_tot = total_mil / total_mfr if total_mfr > 0 else None
    tr = tbl.add_row()
    for i in range(9):
        _set_shading(tr.cells[i], 'D4E1F7')
    _cell(tr.cells[0], 'System Total', bold=True, size=9, align='left')
    _cell(tr.cells[5],
          f'{ratio_tot:.2f}×' if ratio_tot else '—', bold=True, size=9)
    _cell(tr.cells[6], _fmt(total_mil, 5), bold=True, size=9)
    _cell(tr.cells[7], _fmt(total_mfr, 5), bold=True, size=9)
    _cell(tr.cells[8],
          f'{pct_total:+.1f}%' if pct_total is not None else '—',
          bold=True, size=9)


# ═══════════════════════════════════════════════════════════════════════════
#  Main entry point
# ═══════════════════════════════════════════════════════════════════════════

def generate_comparison_report(results:      list[dict],
                               conditions:   dict,
                               project_info: dict) -> bytes:
    """Build English dual-analysis Word report. Returns .docx bytes."""
    from mil217.constants import ENVIRONMENTS_HE, QUALITY_LEVELS

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    sec               = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    style           = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ── Meta ────────────────────────────────────────────────────────────────
    env_code   = conditions.get('env_code',    'GF')
    quality    = conditions.get('quality',     'C')
    t_amb      = conditions.get('t_ambient',    40)
    stress     = conditions.get('stress_ratio', 0.5)
    qual_label = QUALITY_LEVELS.get(quality, quality)

    mil_sum    = _compute_summary(results, 'mil217_lambda_total')
    mfr_sum    = _compute_summary(results, 'mfr_lambda_total')
    ratio_sys  = (mil_sum['total_lambda'] / mfr_sum['total_lambda']
                  if mfr_sum['total_lambda'] > 0 else 1.0)
    alts_data  = _get_top5_alternatives(results)

    # ════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    _add_para(doc, 'COMPARATIVE RELIABILITY ANALYSIS REPORT',
              bold=True, size=22, color=(0x1F, 0x3A, 0x6E), align='center')
    _add_para(doc, 'MIL-STD-217F Notice 2  ·  Manufacturer FIT Database',
              size=14, color=(0x44, 0x44, 0x44), align='center')
    doc.add_paragraph()

    for label, val in [
        ('Project',               project_info.get('project_name', '—')),
        ('Document No.',          project_info.get('doc_number',   'REL-001')),
        ('Prepared by',           project_info.get('prepared_by',  '—')),
        ('Revision',              project_info.get('revision',     'A')),
        ('Date',                  date.today().strftime('%d %B %Y')),
        ('Analysis Temperature',  f'{t_amb} °C'),
        ('Environment Code',      env_code),
        ('Quality Level',         f'{quality} – {qual_label}'),
        ('Components Analysed',   str(len(results))),
    ]:
        p  = doc.add_paragraph()
        r1 = p.add_run(f'{label}:  ')
        r1.font.name = 'Arial'; r1.font.size = Pt(11); r1.bold = True
        r2 = p.add_run(val)
        r2.font.name = 'Arial'; r2.font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '0  Table of Contents', level=1)
    for num, title in [
        ('1',  'Introduction'),
        ('2',  'Methodology'),
        ('3',  'Assumptions & Boundary Conditions'),
        ('4',  'References'),
        ('5',  'MIL-STD-217F Notice 2 – Analysis Results'),
        ('6',  'Manufacturer FIT Database – Analysis Results'),
        ('7',  'Comparison Table'),
        ('8',  'Pareto Chart – Top Failure Contributors'),
        ('9',  'Alternative Component Recommendations (Top 5)'),
        ('10', 'Summary & Conclusions'),
    ]:
        _add_para(doc, f'Section {num}  –  {title}', size=11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 – Introduction
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '1  Introduction', level=1)
    _add_para(doc,
        'This report presents a dual-methodology reliability analysis of the electronic '
        'assembly defined by the attached Bill of Materials (BOM). '
        'Two complementary approaches are applied in parallel:', size=11)
    for bullet in [
        'MIL-STD-217F Notice 2 – US Department of Defense standard for reliability '
        'prediction of electronic equipment. Physics-based models apply environment, '
        'quality, and temperature derating factors (π_T, π_Q, π_E) per component type.',

        'Manufacturer FIT Database – Failure-In-Time rates published by component '
        'manufacturers (TI, Analog Devices, Vishay, Murata, Micron, Semtech, ON Semi, '
        'Diodes Inc, Microchip, Renesas, NXP, ST, Lattice, Nexperia …) in their '
        'Qualification Reports and Reliability Handbooks. Temperature correction is '
        'applied via the Arrhenius equation.',
    ]:
        _add_para(doc, f'•  {bullet}', size=11)

    _add_para(doc, '', size=4)
    _add_para(doc,
        'Comparing both methodologies provides a balanced reliability estimate and '
        'highlights components where the two approaches diverge significantly.',
        size=11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 – Methodology
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '2  Methodology', level=1)

    _heading(doc, '2.1  MIL-STD-217F Notice 2', level=2)
    _add_para(doc,
        'General failure-rate equation  (λ_p in failures / 10⁶ hours):', size=11)
    _add_para(doc,
        'λ_p  =  λ_b × π_T × π_Q × π_E × [additional factors by type]',
        bold=True, size=11, color=(0x1F, 0x3A, 0x6E))

    for sym, desc in [
        ('λ_b',       'Base failure rate – specific to component type and construction'),
        ('π_T',       'Temperature factor – Arrhenius-based with component Eₐ'),
        ('π_Q',       'Quality factor – derived from component qualification level'),
        ('π_E',       'Environment factor – mechanical, thermal and humidity stresses'),
        ('π_R / π_CV','Additional factors for resistance, capacitance, or voltage stress'),
    ]:
        _add_para(doc, f'    •  {sym}  –  {desc}', size=10)

    _add_para(doc, '', size=4)
    for comp, sect, formula in [
        ('Integrated Circuits',      '§5.1', '(C₁·π_T + C₂·π_E)·π_Q·π_L'),
        ('Transistors',              '§6',   'λ_b·π_T·π_A·π_S·π_Q·π_E'),
        ('Diodes',                   '§7',   'λ_b·π_T·π_S·π_C·π_Q·π_E'),
        ('Film Resistors',           '§9',   'λ_b·π_T·π_R·π_Q·π_E'),
        ('Ceramic Capacitors',       '§10',  'λ_b·π_CV·π_T·π_Q·π_E'),
        ('Inductors / Transformers', '§11',  'λ_b·π_T·π_Q·π_E'),
        ('Connectors',               '§15',  'λ_b·π_E·π_Q·π_P'),
        ('Crystals / Oscillators',   '§18',  'λ_b·π_T·π_Q·π_E'),
    ]:
        _add_para(doc, f'    •  {comp}  ({sect}):  {formula}', size=10)

    doc.add_paragraph()

    _heading(doc, '2.2  Manufacturer FIT Database – Arrhenius Correction', level=2)
    _add_para(doc,
        'FIT rates (Failures In Time = failures / 10⁹ hours) are taken from published '
        'manufacturer qualification reports at 25 °C. '
        'Operating-temperature correction uses the Arrhenius equation:', size=11)
    _add_para(doc,
        'FIT(T) = FIT₂₅ × exp[ Eₐ / k_B × ( 1/T_ref − 1/T_use ) ]',
        bold=True, size=11, color=(0x1F, 0x3A, 0x6E))

    for sym, desc in [
        ('FIT₂₅', 'Published FIT rate at 25 °C  (from manufacturer qualification report)'),
        ('Eₐ',    'Activation energy [eV] – failure-mechanism characteristic (JEDEC JEP122H)'),
        ('k_B',   'Boltzmann constant: 8.617 × 10⁻⁵ eV/K'),
        ('T_ref', '298.15 K  (25 °C)'),
        ('T_use', 'T_ambient + 273.15 K'),
    ]:
        _add_para(doc, f'    •  {sym}  –  {desc}', size=10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 – Assumptions
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '3  Assumptions & Boundary Conditions', level=1)
    for a in [
        f'Operating environment: {env_code}',
        f'Maximum ambient temperature: {t_amb} °C',
        f'Voltage stress ratio (V / V_rated): {stress:.0%}',
        f'Component quality level (MIL-217F): {quality} – {qual_label}',
        'Learning factor π_L = 0.5  (production period assumed > 2 years)',
        'Constant Failure Rate (CFR) region assumed – infant-mortality and wear-out excluded',
        'Junction temperature rise above ambient is not modelled (conservative estimate)',
        'Components not matching a specific MIL-217F model are classified as generic digital ICs',
        'Manufacturer FIT rates are at 25 °C standard test conditions with Arrhenius correction applied',
        'Alternative component FIT rates (Section 9) are at 25 °C and representative; '
        'consult manufacturer qualification data for exact values',
    ]:
        _add_para(doc, f'•  {a}', size=11)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 – References
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '4  References', level=1)
    refs = [
        ('[1]',  'MIL-STD-217F Notice 2, "Reliability Prediction of Electronic Equipment," DoD, 1995.', ''),
        ('[2]',  'MIL-HDBK-338B, "Electronic Reliability Design Handbook," 1998.', ''),
        ('[3]',  'JEDEC JEP122H, "Failure Mechanisms and Models for Semiconductor Devices," 2016.', ''),
        ('[4]',  'Vishay – Resistor / Discrete Semiconductor Reliability Data.',
                 'https://www.vishay.com/search/?searchChoice=part&query='),
        ('[5]',  'KEMET / Murata – MLCC Reliability Reports.', ''),
        ('[6]',  'Analog Devices (incl. Linear Technology / Maxim) – Reliability Data.',
                 'https://www.analog.com/en/about-adi/quality-reliability/reliability-data/wafer-fabrication-data.html'),
        ('[7]',  'Texas Instruments – Quality & Reliability / FIT Estimator.',
                 'https://www.ti.com/quality/docs/estimator.tsp'),
        ('[8]',  'ON Semiconductor – Reliability Data.',
                 'https://www.onsemi.com/PowerSolutions/reliability.do'),
        ('[9]',  'Semtech – FIT Rate Estimator.',
                 'https://www.semtech.com/quality/reliability'),
        ('[10]', 'Diodes Inc – MTBF / FIT Estimator.',
                 'https://www.diodes.com/quality/mtbffit-estimator'),
        ('[11]', 'Microchip Technology – Reliability Report.',
                 'https://www.microchip.com/reliabilityreport/#/'),
        ('[12]', 'MCC Semiconductor – Reliability Data.',
                 'https://www.mccsemi.com/ReliabilityData/'),
        ('[13]', 'Central Semiconductor – FIT Rate Data.',
                 'https://www.centralsemi.com/reliability-data#reliability-FIT-rate'),
        ('[14]', 'Renesas Electronics – Reliability / Document Search.',
                 'https://www.renesas.com/en/support/document-search'),
        ('[15]', 'NXP Semiconductors – Reliability Data.',
                 'https://www.nxp.com/company/about-nxp/quality/reliability-data:RELIABILITY-DATA'),
        ('[16]', 'STMicroelectronics – Quality & Reliability.',
                 'https://www.st.com/content/st_com/en/support/quality-reliability.html'),
        ('[17]', 'Lattice Semiconductor – Quality & Reliability.',
                 'https://www.latticesemi.com/en/Support/QualityAndReliability'),
        ('[18]', 'Nexperia – Reliability Data.',
                 'https://www.nexperia.com/quality/reliability/'),
        ('[19]', 'ISSI – Quality & Reliability.',
                 'https://www.issi.com/US/quality-reliability.shtml'),
        ('[20]', 'Würth Elektronik – Component Quality.',
                 'https://www.we-online.com/en/components/products/quality'),
        ('[21]', 'Reliability Analytics Toolkit (general reference).',
                 'https://reliabilityanalyticstoolkit.appspot.com'),
        ('[22]', 'Micron Technology – Memory Product Reliability Report.', ''),
        ('[23]', 'Infineon Technologies – FRAM / Power Reliability Data.', ''),
        ('[24]', 'CTS Corporation – Oscillator / TCXO Qualification Report.', ''),
        ('[25]', 'Intel / Altera – Agilex Series Reliability Qualification.', ''),
    ]
    for num, text, url in refs:
        para = doc.add_paragraph()
        r1   = para.add_run(f'{num}  {text}')
        r1.font.name = 'Arial'; r1.font.size = Pt(10)
        if url:
            para.add_run('  ').font.size = Pt(10)
            _para_hyperlink(para, url, url, size=9)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 – MIL-217F results
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '5  MIL-STD-217F Notice 2 – Analysis Results', level=1)
    _add_para(doc,
        f'Environment: {env_code}  |  Temperature: {t_amb} °C  |  '
        f'Quality: {quality}  |  Stress ratio: {stress:.0%}',
        size=10, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()

    _build_mil217_table(doc, results)

    _add_para(doc, '', size=4)
    _add_para(doc,
        f'System total  λ_sys (MIL-217F) = {mil_sum["total_lambda"]:.4f} failures/10⁶ hr  '
        f'({mil_sum["fit_rate"]:.1f} FIT)   →   '
        f'MTBF = {mil_sum["mtbf_hours"]:,.0f} hr  ({mil_sum["mtbf_years"]:.1f} years)',
        bold=True, size=11, color=(0x1F, 0x3A, 0x6E))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 – Manufacturer data results
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '6  Manufacturer FIT Database – Analysis Results', level=1)
    _add_para(doc,
        f'Operating temperature: {t_amb} °C  |  Temperature model: Arrhenius',
        size=10, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()

    _build_mfr_table(doc, results)

    _add_para(doc, '', size=4)
    _add_para(doc,
        f'System total  λ_sys (Manufacturer) = {mfr_sum["total_lambda"]:.4f} failures/10⁶ hr  '
        f'({mfr_sum["fit_rate"]:.1f} FIT)   →   '
        f'MTBF = {mfr_sum["mtbf_hours"]:,.0f} hr  ({mfr_sum["mtbf_years"]:.1f} years)',
        bold=True, size=11, color=(0x19, 0x87, 0x54))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 – Comparison table
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '7  Comparison Table', level=1)
    _add_para(doc,
        'The table below compares both analyses per component. '
        'Ratio > 1 indicates MIL-STD-217F is more conservative than the manufacturer data. '
        'Orange shading = ratio > 5× (large discrepancy); '
        'yellow = 2–5×; white < 2× (consistent).',
        size=11)
    doc.add_paragraph()

    _build_comparison_table(doc, results)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 8 – Pareto chart
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '8  Pareto Chart – Top Failure Rate Contributors', level=1)
    _add_para(doc,
        'The chart below ranks components by total failure rate contribution '
        '(λ_p × quantity) using the manufacturer FIT database. '
        'The cumulative line identifies the vital few components that drive the '
        'majority of system failures (Pareto / 80–20 rule). '
        'The top 5 bars (red) are the candidates addressed in Section 9.',
        size=11)
    doc.add_paragraph()

    chart_png = _generate_pareto_chart(results)
    if chart_png:
        doc.add_picture(_io.BytesIO(chart_png), width=Inches(6.5))
        _add_para(doc,
            'Figure 8.1 – Pareto Chart: Component Failure-Rate Contribution  '
            '(Manufacturer FIT Data)',
            size=9, italic=True, color=(0x55, 0x55, 0x55), align='center')
    else:
        _add_para(doc,
            '[Pareto chart unavailable – install matplotlib: pip install matplotlib]',
            size=10, italic=True, color=(0x88, 0x88, 0x88))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9 – Alternative component recommendations
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '9  Alternative Component Recommendations – Top 5 Contributors', level=1)
    _add_para(doc,
        'The five components with the highest λ_p_total (manufacturer FIT data) are '
        'analysed below. For each, alternative devices with lower published FIT rates are '
        'suggested. FIT rates for alternatives are at 25 °C before Arrhenius correction; '
        'the projected MTBF improvement assumes a direct drop-in replacement.',
        size=11)
    doc.add_paragraph()

    for i, item in enumerate(alts_data['items'], start=1):
        comp = item['component']
        _heading(doc,
                 f'9.{i}  {comp["ref_des"]}  –  '
                 f'{comp.get("mfr_category", "?")}', level=2)

        # Current component summary
        p = doc.add_paragraph()
        p.add_run(
            f'Part: {comp["part_number"]}  |  '
            f'Manufacturer: {comp["manufacturer"]}  |  '
            f'Qty: {comp["quantity"]}  |  '
            f'Description: {comp["description"][:60]}'
        ).font.size = Pt(10)

        p2 = doc.add_paragraph()
        r2 = p2.add_run(
            f'Current  λ_p_total = {_fmt(item["cur_lt"], 5)} failures/10⁶ hr  '
            f'(FIT@25°C = {_fmt(comp.get("mfr_fit_25c", 0), 1)})'
        )
        r2.font.size = Pt(10)

        p3 = doc.add_paragraph()
        rb = p3.add_run('Root cause of high FIT:  ')
        rb.font.size = Pt(10); rb.bold = True
        p3.add_run(item['reason']).font.size = Pt(10)

        if item['alternatives']:
            ph = doc.add_paragraph()
            rh = ph.add_run('Recommended alternatives:')
            rh.font.size = Pt(10); rh.bold = True

            # Alternatives mini-table
            at = doc.add_table(rows=1, cols=4)
            at.style     = 'Table Grid'
            at.alignment = WD_TABLE_ALIGNMENT.CENTER
            for hi, hdr in enumerate(['Alternative Part', 'FIT@25°C', 'Notes', 'Source']):
                _set_shading(at.rows[0].cells[hi], '1F3A6E')
                _cell(at.rows[0].cells[hi], hdr, bold=True, size=9,
                      color_rgb=(255, 255, 255))
            for ci, w in zip(range(4), [5.0, 1.8, 6.5, 3.0]):
                for c in at.columns[ci].cells:
                    c.width = Cm(w)

            for alt in item['alternatives']:
                ar = at.add_row()
                _cell(ar.cells[0], alt['part'],  size=9, align='left')
                _cell(ar.cells[1], str(alt['fit_25c']), size=9)
                _cell(ar.cells[2], alt['notes'], size=8, align='left')
                _cell_link(ar.cells[3], 'Reliability Data', alt.get('url', ''), size=8)

        if item['best_alt']:
            ba     = item['best_alt']
            new_lp = ba['fit_25c'] / 1000.0
            new_lt = new_lp * comp['quantity']
            pct_imp = ((item['cur_lt'] - new_lt) / item['cur_lt'] * 100
                       if item['cur_lt'] > 0 else 0)
            p5 = doc.add_paragraph()
            p5.add_run(
                f'Best-case gain  (using  {ba["part"]}):  '
                f'λ_p_total  {_fmt(item["cur_lt"], 5)}  →  {_fmt(new_lt, 5)} '
                f'failures/10⁶ hr   (↓ {pct_imp:.0f}%)'
            ).font.size = Pt(10)
        else:
            _add_para(doc,
                'No specific alternative identified for this category in the database.',
                size=10, italic=True)

        doc.add_paragraph()

    # ── Aggregated MTBF improvement ─────────────────────────────────────────
    _heading(doc, '9.6  Projected System MTBF After Top-5 Replacement', level=2)
    _add_para(doc,
        'The table below shows the projected system MTBF assuming all five recommended '
        'best-alternative components are adopted simultaneously.',
        size=11)
    doc.add_paragraph()

    sum_tbl = doc.add_table(rows=5, cols=3)
    sum_tbl.style     = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for hi, hdr in enumerate(['Parameter', 'Current (as-designed)', 'After Top-5 Replacement']):
        _set_shading(sum_tbl.rows[0].cells[hi], '1F3A6E')
        _cell(sum_tbl.rows[0].cells[hi], hdr, bold=True, size=10, color_rgb=(255, 255, 255))

    imp_rows = [
        ('λ_sys  (failures/10⁶ hr)',
         f'{alts_data["current_total"]:.4f}',
         f'{alts_data["new_total"]:.4f}'),
        ('System FIT  (failures/10⁹ hr)',
         f'{alts_data["current_total"] * 1000:.1f}',
         f'{alts_data["new_total"] * 1000:.1f}'),
        ('MTBF  (hours)',
         f'{alts_data["current_mtbf_h"]:,.0f}',
         f'{alts_data["new_mtbf_h"]:,.0f}'),
        ('MTBF  (years)',
         f'{alts_data["current_mtbf_y"]:.1f}',
         f'{alts_data["new_mtbf_y"]:.1f}'),
    ]
    for ri, (lbl, cur, new) in enumerate(imp_rows):
        row = sum_tbl.rows[ri + 1]
        bg  = 'EEF2F7' if ri % 2 == 0 else 'FFFFFF'
        _set_shading(row.cells[0], bg)
        _cell(row.cells[0], lbl, bold=True, size=10, align='left')
        _set_shading(row.cells[1], bg)
        _cell(row.cells[1], cur, size=10)
        _set_shading(row.cells[2], '90EE90' if ri >= 2 else bg)
        _cell(row.cells[2], new, size=10,
              color_rgb=(0, 100, 0) if ri >= 2 else None)

    doc.add_paragraph()
    _add_para(doc,
        f'Replacing the top-5 highest-FIT components with recommended alternatives '
        f'is projected to improve system MTBF by  {alts_data["improvement_factor"]:.2f}×  '
        f'({alts_data["current_mtbf_y"]:.1f} yr  →  {alts_data["new_mtbf_y"]:.1f} yr).',
        bold=True, size=11, color=(0, 100, 0))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10 – Summary & Conclusions
    # ════════════════════════════════════════════════════════════════════════
    _heading(doc, '10  Summary & Conclusions', level=1)

    _heading(doc, '10.1  KPI Summary', level=2)
    kpi_tbl = doc.add_table(rows=5, cols=3)
    kpi_tbl.style     = 'Table Grid'
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for hi, hdr in enumerate(['Parameter', 'MIL-STD-217F', 'Manufacturer FIT']):
        _set_shading(kpi_tbl.rows[0].cells[hi], '1F3A6E')
        _cell(kpi_tbl.rows[0].cells[hi], hdr, bold=True, size=10,
              color_rgb=(255, 255, 255))

    for ri, (lbl, mv, fv) in enumerate([
        ('λ_sys  (failures/10⁶ hr)',      f'{mil_sum["total_lambda"]:.4f}', f'{mfr_sum["total_lambda"]:.4f}'),
        ('System FIT  (failures/10⁹ hr)', f'{mil_sum["fit_rate"]:.1f}',     f'{mfr_sum["fit_rate"]:.1f}'),
        ('MTBF  (hours)',                  f'{mil_sum["mtbf_hours"]:,.0f}',  f'{mfr_sum["mtbf_hours"]:,.0f}'),
        ('MTBF  (years)',                  f'{mil_sum["mtbf_years"]:.1f}',   f'{mfr_sum["mtbf_years"]:.1f}'),
    ]):
        row = kpi_tbl.rows[ri + 1]
        bg  = 'EEF2F7' if ri % 2 == 0 else 'FFFFFF'
        for ci in range(3):
            _set_shading(row.cells[ci], bg)
        _cell(row.cells[0], lbl, bold=True, size=10, align='left')
        _cell(row.cells[1], mv, size=10)
        _cell(row.cells[2], fv, size=10)

    doc.add_paragraph()

    _heading(doc, '10.2  Analysis Interpretation', level=2)
    if ratio_sys > 5:
        conservatism = 'very conservative'
    elif ratio_sys > 2:
        conservatism = 'conservative'
    elif ratio_sys > 1.2:
        conservatism = 'slightly conservative'
    elif ratio_sys > 0.8:
        conservatism = 'consistent'
    else:
        conservatism = 'relatively optimistic'

    _add_para(doc,
        f'MIL-217F / Manufacturer FIT ratio: {ratio_sys:.2f}×  –  '
        f'MIL-STD-217F is {conservatism} relative to manufacturer data.',
        bold=True, size=11)
    doc.add_paragraph()

    mtbf_cons   = min(mil_sum['mtbf_hours'], mfr_sum['mtbf_hours'])
    mtbf_cons_y = mtbf_cons / 8760
    if mtbf_cons > 100_000:
        assessment = (
            f'Both methodologies indicate high system reliability. '
            f'The most conservative MTBF estimate is {mtbf_cons:,.0f} hr '
            f'({mtbf_cons_y:.1f} years), meeting typical industrial / '
            f'automotive requirements.'
        )
    elif mtbf_cons > 20_000:
        assessment = (
            f'The most conservative MTBF is {mtbf_cons:,.0f} hr ({mtbf_cons_y:.1f} yr) – '
            f'adequate for commercial / industrial equipment. '
            f'Review the top-5 failure contributors (Section 9) for targeted improvements.'
        )
    else:
        assessment = (
            f'The most conservative MTBF is {mtbf_cons:,.0f} hr ({mtbf_cons_y:.1f} yr). '
            f'Design improvements are warranted: component quality upgrade, thermal derating, '
            f'and / or substitution of high-FIT parts (see Section 9).'
        )
    _add_para(doc, assessment, size=11)
    doc.add_paragraph()

    _heading(doc, '10.3  Recommendations', level=2)
    for rec in [
        f'Replace top-5 high-FIT components (Section 9) with recommended alternatives – '
        f'projected MTBF improvement: {alts_data["improvement_factor"]:.1f}× '
        f'({alts_data["current_mtbf_y"]:.1f} yr → {alts_data["new_mtbf_y"]:.1f} yr).',

        'Apply thermal management to components with π_T > 3.0; every 10 °C reduction '
        'approximately halves the failure rate for silicon ICs.',

        f'Upgrade critical components from Quality Level {quality} to Level B-2 or higher '
        f'(MIL-217F) to reduce π_Q.',

        'Perform Burn-In screening on high-complexity ICs to eliminate infant-mortality failures.',

        'Re-run this analysis after any BOM change, environment change, or quality-level update.',

        'Components with MIL-217F / Manufacturer ratio > 5× should be reviewed individually – '
        'the MIL model may be overly conservative for modern process nodes.',
    ]:
        _add_para(doc, f'•  {rec}', size=11)

    doc.add_paragraph()
    _add_para(doc, '─' * 80, size=9)
    _add_para(doc,
        'This report was generated automatically. '
        'Verify all input assumptions before using results for certification or design sign-off.',
        size=9, italic=True, color=(0x88, 0x88, 0x88))

    # ── Save ────────────────────────────────────────────────────────────────
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
